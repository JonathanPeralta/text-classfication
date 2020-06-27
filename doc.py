# import zipfile, re
# filename = "sobras.docx"


# docx = zipfile.ZipFile(filename)
# content = docx.read('Word/document.xml').decode('utf-8')
# cleaned = re.sub('<(.|\n)*?>','',content)
# print(cleaned)
import docx


doc = docx.Document('sobras.docx')
# print(len(doc.paragraphs))

var = ""
for i in doc.paragraphs:
    if i.text != "":
        var += i.text
        # print(i.text)
# print(doc.paragraphs[0].text)

print(var)