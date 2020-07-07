from django.shortcuts import render
# from . forms import MyForm
import json
import logging
import requests

from bs4 import BeautifulSoup
from rest_framework import viewsets
from rest_framework.decorators import api_view
from django.core import serializers
from rest_framework.response import Response
from rest_framework import status
from django.http import JsonResponse
from rest_framework.parsers import JSONParser
# from . models import approvals
from . models import Document
# from . serializers import approvalsSerializers
from . serializer import DocumentSerializar
# import pickle
# from sklearn.externals import joblib
# import json
# import numpy as np
# from sklearn import preprocessing
# import pandas as pd
import sklearn.datasets as skd
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.feature_extraction.text import TfidfTransformer
from sklearn.naive_bayes import MultinomialNB

import nltk
import inflect
import re
#nltk.download('stopwords')
from nltk.corpus import stopwords
stop_words = set(stopwords.words('spanish'))

import docx
import tika
from tika import parser
from tika import detector

logger = logging.getLogger(__name__)



class DocumentView(viewsets.ModelViewSet):
	queryset = Document.objects.all()
	serializer_class = DocumentSerializar

@api_view(["POST"])
def prueba(request):
    try:
        mydata=request.FILES
        d = list(mydata.values())
        document = Document(archive=d[0])
        document.save()

        lastdocument = Document.objects.get(pk=document.id)

        patharchive = str(lastdocument.archive)

        archivename = patharchive.split('/')[1]
        archivetype = archivename.split('.')[1]
        
        if archivetype=="pdf":
            text = textpdf(patharchive)
        else:
            text = textword(patharchive)


        # words = nltk.word_tokenize(text)
        # words = normalize(words)


        category = proccess(text)

        # print(words)
        # print(category)
        keywords = keywords_function(category)

      
        
        URL = 'https://pe.indeed.com/jobs?q='+category+'&l='
        INDEED = 'https://pe.indeed.com'

        response = requests.get(URL)
        html_soup = BeautifulSoup(response.text, 'html.parser')

        job_containers = html_soup.find_all('div', class_ = 'jobsearch-SerpJobCard')

        jobs = []
        id = 1
        for etq in job_containers:
            
            anchordiv = etq.h2
            anchor = anchordiv.find('a', href=True)
            link = anchor['href']

            footer = etq.find('div',class_="jobsearch-SerpJobCard-footer")
            span = footer.find('span',class_="date")
            date = span.text

            title = str(etq.h2.a.text).strip()

            companydiv = etq.find('div', class_ = 'sjcl')
            # print(em)
            companydiv.find('div',class_="company")
            company = str(companydiv.span.text).strip()

            companydiv.find("div",class_="location")
            precity = str(companydiv.text).strip()
            city = precity.split('\n\n\n')[1]
            city = city.strip()
            # print(city)


            summary = etq.find("div",class_="summary")
            # print(summary.ul)
            uls = summary.ul

            lis=[]
            idd = 1
            for li in uls.find_all("li", recursive=True): 
                dic = {'id':"idd"+str(id)+str(idd),'li':li.text}
                lis.append(dic)
                idd=idd+1
            # obj =  "{ 'title':'"+title+"','company':'"+company+"','city':'"+city+"','description':'"+lis+"'}"
            obj = {'id':"id"+str(id),'date':date,'link':INDEED+link,'title':title,'company':company,'city':city,'description':lis}

            jobs.append(obj)
            id=id+1

        return JsonResponse({'lista':category,'res':'ok',"jobs":jobs}, safe=False)

    except ValueError as e:
        return Response(e.args[0], status.HTTP_400_BAD_REQUEST)







def proccess(document):

    categories = ['Administrador de sistemas y Devops', 
                  'Analista financiero',
                  'Asociado consultor', 
                  'Asociado de proyecto',
                  'Desarrollador de aplicaciones móviles',
                  'Desarrollador de Javascript',
                  'Desasrrollador Full Stack',
                  'Desarrollador Ruby On Rails',
                  'Desasrrollador web',
                  'Desarrollador de Java',
                  'Desarrollador web PHP',
                  'Desarrollo web python',
                  'Diseñador gráfico',
                  'Editor Asociado',
                  'Especialista en datos',
                  'Generalista de Recursos Humanos Asociado',
                  'Ingeniero de diseño de hardware',
                  'Ingeniero de pruebas de calidad de software',
                  'Ingeniero de software',
                  'Ingeniero de ventas',
                  'marketing digital',
                  'Oficial de recaudación de fondos'];

    news_train = skd.load_files('files/nuevos/train/', categories= categories, encoding= 'ISO-8859-1')
    news_test = skd.load_files('files/nuevos/test/',categories= categories, encoding= 'ISO-8859-1')

    # print(news_train.target_names)

    count_vect = CountVectorizer()
    x_train_tf = count_vect.fit_transform(news_train.data)
    x_train_tf.shape

    tfidf_transformer = TfidfTransformer()
    x_train_tfidf = tfidf_transformer.fit_transform(x_train_tf)
    # print(x_train_tfidf)

    clf = MultinomialNB().fit(x_train_tfidf,news_train.target)

    x_test_tf = count_vect.transform(news_test.data)
    x_test_tfidf =  tfidf_transformer.fit_transform(x_test_tf)
    # predicted = clf.predict(x_test_tfidf)

    
    docs_news = [document]
    x_new_counts = count_vect.transform(docs_news)
    x_new_tfidf =  tfidf_transformer.transform(x_new_counts)
    predicted = clf.predict(x_new_tfidf)

    for x in predicted:
        print(news_train.target_names[x])
        # print(x)
    return news_train.target_names[x]


def keywords_function(category):

        switcher = {
            "Administrador de sistemas y Devops": "",
            "Analista financiero": "",
            "Asociado consultor": "",
            "Asociado de proyecto": "",
            "Desarrollador de aplicaciones móviles": "",
            "Desarrollador de Javascript": "",
            "Desasrrollador Full Stack": "",
            "Desarrollador Ruby On Rails": "",
            "Desasrrollador web": "",
            "Desarrollador de Java": "",
            "Desarrollador web PHP": "",
            "Desarrollo web python": "",
            "Diseñador gráfico":"UX/UI/frontend/illustrator/photoshop/figma/diseño gráfico/diseñador gráfico",
            "Editor Asociado": "",
            "Especialista en datos": "",
            "Generalista de Recursos Humanos Asociado": "",
            "Ingeniero de diseño de hardware": "",
            "Ingeniero de pruebas de calidad de software":"",
            "Ingeniero de ventas": "",
            "marketing digital": "",
            "Ingeniero de software": "",
            "Oficial de recaudación de fondos":"",

        }
        # Get the function from switcher dictionary
        func = switcher.get(category, lambda: "Invalid value")
        return func




def textword(document):
    
    doc = docx.Document(document)
    text = ""
    for i in doc.paragraphs:
        if i.text != "":
            text += i.text

    return text




def textpdf(document):

    tika.initVM()
    parsed = parser.from_file(document.encode('utf-8'))
    return parsed["content"]




# print(words)


def remove_non_ascii(words):
    """Remove non-ASCII characters from list of tokenized words"""
    new_words = []
    for word in words:
        new_word = unicodedata.normalize('NFKD', word).encode('ascii', 'ignore').decode('utf-8', 'ignore')
        new_words.append(new_word)
    return new_words

def to_lowercase(words):
    """Convert all characters to lowercase from list of tokenized words"""
    new_words = []
    for word in words:
        new_word = word.lower()
        new_words.append(new_word)
    return new_words

def remove_punctuation(words):
    """Remove punctuation from list of tokenized words"""
    new_words = []
    for word in words:
        new_word = re.sub(r'[^\w\s]', '', word)
        if new_word != '':
            new_words.append(new_word)
    return new_words


def replace_numbers(words):
    """Replace all interger occurrences in list of tokenized words with textual representation"""
    p = inflect.engine()
    new_words = []
    for word in words:
        if word.isdigit():
            new_word = p.number_to_words(word)
            new_words.append(new_word)
        else:
            new_words.append(word)
    return new_words

def remove_stopwords(words):
    """Remove stop words from list of tokenized words"""
    new_words = []
    for word in words:
        if word not in stopwords.words('english'):
            new_words.append(word)
    return new_words

def stem_words(words):
    """Stem words in list of tokenized words"""
    stemmer = LancasterStemmer()
    stems = []
    for word in words:
        stem = stemmer.stem(word)
        stems.append(stem)
    return stems

def lemmatize_verbs(words):
    """Lemmatize verbs in list of tokenized words"""
    lemmatizer = WordNetLemmatizer()
    lemmas = []
    for word in words:
        lemma = lemmatizer.lemmatize(word, pos='v')
        lemmas.append(lemma)
    return lemmas

def normalize(words):
    # words = remove_non_ascii(words)
    words = to_lowercase(words)
    words = remove_punctuation(words)
    words = replace_numbers(words)
    words = remove_stopwords(words)
    return words

# words = normalize(words)


class Object(): 
    def __init__(self, title):
        self.title = title
    
    def __str__(self):
        json_="{title:"+self.nombre+"}"
        return json_